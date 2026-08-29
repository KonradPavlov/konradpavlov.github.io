"""
Pav.export — Rendu et exportation "publication-ready"
Auteur : SALLAN Konrad Pavlov
 et exportation "publication-ready" des tableaux :
Markdown propre (GitHub-flavored), Microsoft Word (.docx) et HTML interactif.
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional, Union

import pandas as pd
from tabulate import tabulate


# ---------------------------------------------------------------------------
# Export Markdown
# ---------------------------------------------------------------------------
def to_markdown(df: pd.DataFrame, title: Optional[str] = None) -> str:
    """
    Rendu Markdown propre (GitHub-flavored) en préservant strictement
    les zéros décimaux des p-values et des pourcentages (ex: '1.00', '0.010').
    """
    lines = []
    if title:
        lines.append(f"## {title}")
        lines.append("")
    # disable_numparse=True empêche tabulate de convertir les chaînes '1.000' en float 1
    md_table = tabulate(df, headers="keys", tablefmt="pipe", showindex=False, disable_numparse=True)
    lines.append(md_table)
    return "\n".join(lines)


def save_markdown(df: pd.DataFrame, path: Union[str, Path], title: Optional[str] = None) -> None:
    """Enregistre le tableau au format Markdown."""
    Path(path).write_text(to_markdown(df, title=title), encoding="utf-8")


# ---------------------------------------------------------------------------
# Export HTML
# ---------------------------------------------------------------------------
def to_html(
    df: pd.DataFrame,
    path: Optional[Union[str, Path]] = None,
    title: Optional[str] = None,
    footnote: Optional[str] = None,
) -> str:
    """
    Génère un tableau HTML élégant avec style CSS médical / publication-ready intégré.
    """
    html_lines = [
        "<!DOCTYPE html>",
        "<html lang='fr'>",
        "<head>",
        "<meta charset='utf-8'>",
        "<style>",
        "body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; margin: 20px; color: #222; }",
        "h2 { color: #1e293b; font-size: 1.25rem; margin-bottom: 0.75rem; }",
        "table.Pav-table { border-collapse: collapse; width: 100%; max-width: 1000px; margin-bottom: 1rem; font-size: 0.9rem; }",
        "table.Pav-table th { background-color: #f1f5f9; color: #0f172a; font-weight: 600; text-align: left; padding: 8px 12px; border-top: 2px solid #0f172a; border-bottom: 1.5px solid #0f172a; }",
        "table.Pav-table td { padding: 6px 12px; border-bottom: 1px solid #e2e8f0; vertical-align: top; }",
        "table.Pav-table tr.header-row { background-color: #f8fafc; font-weight: 600; }",
        "table.Pav-table tr:hover { background-color: #f8fafc; }",
        ".footnote { font-size: 0.8rem; color: #64748b; font-style: italic; margin-top: 0.5rem; }",
        "</style>",
        "</head>",
        "<body>",
    ]

    if title:
        html_lines.append(f"<h2>{title}</h2>")

    html_lines.append("<table class='Pav-table'>")
    html_lines.append("<thead><tr>")
    for col in df.columns:
        html_lines.append(f"<th>{col}</th>")
    html_lines.append("</tr></thead>")
    html_lines.append("<tbody>")

    for _, row in df.iterrows():
        first_val = str(row.iloc[0]).strip()
        is_header_row = (first_val.startswith("**") and first_val.endswith(":**")) or (first_val.startswith("**N ="))
        tr_class = " class='header-row'" if is_header_row else ""
        html_lines.append(f"<tr{tr_class}>")
        for val in row:
            val_str = "" if pd.isna(val) else str(val)
            # Remplacer les marqueurs Markdown **texte** par <strong>texte</strong>
            while "**" in val_str:
                val_str = val_str.replace("**", "<strong>", 1).replace("**", "</strong>", 1)
            html_lines.append(f"<td>{val_str}</td>")
        html_lines.append("</tr>")

    html_lines.append("</tbody></table>")

    if footnote:
        html_lines.append(f"<div class='footnote'>{footnote}</div>")

    html_lines.append("</body></html>")
    full_html = "\n".join(html_lines)

    if path is not None:
        Path(path).write_text(full_html, encoding="utf-8")

    return full_html


# ---------------------------------------------------------------------------
# Export Word (.docx)
# ---------------------------------------------------------------------------
def _parse_bold_segments(text: str):
    segments = []
    remaining = text
    while "**" in remaining:
        before, _, rest = remaining.partition("**")
        if before:
            segments.append((before, False))
        bold_text, _, remaining = rest.partition("**")
        segments.append((bold_text, True))
    if remaining:
        segments.append((remaining, False))
    return segments or [("", False)]


def to_docx(
    df: pd.DataFrame,
    path: Union[str, Path],
    title: Optional[str] = None,
    footnote: Optional[str] = None,
) -> None:
    """
    Exporte le tableau en document Word (.docx) avec mise en forme professionnelle
    (en-têtes ombrés, bordures fines, colonnes adaptées).
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    if title:
        doc.add_heading(title, level=2)

    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Table Grid"

    def _set_cell_shading(cell, hex_color: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _write_cell(cell, text: str, bold_all: bool = False, header: bool = False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for seg_text, seg_bold in _parse_bold_segments(str(text)):
            if not seg_text:
                continue
            run = p.add_run(seg_text)
            run.bold = bold_all or seg_bold or header
            run.font.size = Pt(9 if not header else 10)
            if header:
                run.font.color.rgb = RGBColor(0, 0, 0)

    # En-tête
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        _write_cell(cell, str(col), header=True)
        _set_cell_shading(cell, "D9D9D9")

    # Corps
    for i in range(n_rows):
        first_val = str(df.iloc[i, 0]).strip()
        row_is_header = (first_val.startswith("**") and first_val.endswith(":**")) or (first_val.startswith("**N ="))
        for j in range(n_cols):
            cell = table.rows[i + 1].cells[j]
            val = df.iloc[i, j]
            _write_cell(cell, "" if pd.isna(val) else str(val))
            if row_is_header:
                _set_cell_shading(cell, "F2F2F2")

    if footnote:
        p = doc.add_paragraph()
        run = p.add_run(footnote)
        run.italic = True
        run.font.size = Pt(8)

    doc.save(str(path))
